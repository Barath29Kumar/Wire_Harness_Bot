import csv
import os
import uuid
import base64
from io import BytesIO
from PIL import Image
import chromadb
from chromadb.config import Settings

chroma_client = chromadb.PersistentClient(path="./chroma_db", settings=Settings(allow_reset=True))
chroma_client = chromadb.Client(Settings(allow_reset=True))

from sentence_transformers import SentenceTransformer, CrossEncoder
import pandas as pd
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from PyPDF2 import PdfReader
import http
import json
from docx import Document as docxDocument
import re
import os
from dotenv import load_dotenv
from pathlib import Path
import streamlit as st
 
# Specify the path to your .env file
env_path = Path('../.env')  # Adjust this path as needed
 
# Load the .env file
load_dotenv(dotenv_path=env_path)
 
# Now you can access your environment variables
api_key = os.getenv('API_KEY')
# ChromaDB client setup
def get_chroma_client():
    PERSIST_DIRECTORY = os.path.join(os.getcwd(), "chromadb_data")
    os.makedirs(PERSIST_DIRECTORY, exist_ok=True)
    return chromadb.PersistentClient(path=PERSIST_DIRECTORY)
 
# Sentence transformers setup
model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
query_encoder = SentenceTransformer('sentence-transformers/msmarco-distilbert-base-v4')
passage_encoder = SentenceTransformer('sentence-transformers/msmarco-distilbert-base-v4')
cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')
 
# def create_collection(collection_name, chunks):
#     chroma_client = get_chroma_client()
#     try:
#         collection = chroma_client.get_collection(name=collection_name)
#     except chromadb.errors.InvalidCollectionException:
#         collection = chroma_client.create_collection(name=collection_name)
   
#     for chunk in chunks:
#         embedding = passage_encoder.encode(chunk).tolist()
#         unique_id = str(uuid.uuid4())
#         collection.add(
#             documents=[chunk],
#             embeddings=[embedding],
#             ids=[unique_id]
#         )
 
def read_csv_file(file_path):
    text = ""
    encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
    for encoding in encodings:
        try:
            with open(file_path, mode='r', newline='', encoding=encoding) as file:
                csv_reader = csv.reader(file)
                for row in csv_reader:
                    text += ", ".join(row) + "\n"
            break
        except UnicodeDecodeError:
            print(f"Encoding error with {encoding}. Trying next encoding...")
        except Exception as e:
            print(f"Error reading CSV file {file_path} with encoding {encoding}: {e}")
            break
    return text
 
def extract_tables_from_docx(file_path):
    doc = docxDocument(file_path)
    all_tables_data = []
   
    for i, table in enumerate(doc.tables):
        table_data = f"Table {i+1}:\n"
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data += " | ".join(row_data) + "\n"
        all_tables_data.append(table_data)
   
    text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    if text_content:
        all_tables_data.append(text_content)
   
    return all_tables_data
 
def create_collection(collection_name, chunks):
    chroma_client = get_chroma_client()
    try:
        collection = chroma_client.get_collection(name=collection_name)
        print(f"Collection '{collection_name}' already exists")
        chroma_client.delete_collection(name=collection_name)
    except chromadb.errors.InvalidCollectionException:
        pass
   
    collection = chroma_client.create_collection(name=collection_name)
   
    for i, chunk in enumerate(chunks):
        embedding = model.encode(chunk).tolist()
        unique_id = str(uuid.uuid4())
        collection.add(
            documents=[chunk],
            embeddings=[embedding],
            ids=[unique_id]
        )
        if i % 100 == 0:
            print(f"Added {i+1} documents to the collection")
   
    print(f"Collection '{collection_name}' created with {len(chunks)} documents.")
   
    # Verify the collection
    print(f"Collection info: {collection.count()} documents")
 
def chunk_documents(documents, chunk_size=300, overlap=50):
    chunks = []
    for document in documents:
        if isinstance(document, str):
            words = document.split()
            for i in range(0, len(words), chunk_size - overlap):
                chunk = ' '.join(words[i:i + chunk_size])
                chunks.append(chunk)
        elif isinstance(document, dict) and 'text' in document:
            words = document['text'].split()
            for i in range(0, len(words), chunk_size - overlap):
                chunk = ' '.join(words[i:i + chunk_size])
                chunks.append(chunk)
    return chunks
 
def semantic_search(query, collection, top_k=20):
    query_embedding = model.encode(query).tolist()
    results = collection.query(
        query_embeddings=[query_embedding],
        n_results=top_k,
        include=['documents', 'embeddings']
    )
    print(f"Query results: {results}")  # Debug print
   
    if not results['documents'] or 'embeddings' not in results or not results['embeddings']:
        print("No results found or embeddings not returned")
        return [], np.array([])
   
    return results['documents'][0], np.array(results['embeddings'][0])
 
def retrieve_documents(prompt):
    collection_name = "test_files"
    chroma_client = get_chroma_client()
    try:
        collection = chroma_client.get_collection(name=collection_name)
    except chromadb.errors.InvalidCollectionException:
        print(f"Collection '{collection_name}' not found")
        return ""
 
    initial_results, embeddings = semantic_search(prompt, collection)
   
    if not initial_results:
        print("No relevant documents found")
        return ""
   
    if embeddings is None or (isinstance(embeddings, np.ndarray) and embeddings.size == 0):
        print("Embeddings not returned or empty, skipping re-ranking")
        return " ".join(initial_results)
   
    reranked_results = re_rank(prompt, initial_results, embeddings)
   
    context = " ".join(reranked_results)
    return context
 
 
def re_rank(query, documents, embeddings, top_k=5):
    query_embedding = model.encode(query)
    if isinstance(embeddings, list):
        embeddings = np.array(embeddings)
    similarities = cosine_similarity([query_embedding], embeddings)[0]
    ranked_results = sorted(zip(documents, similarities), key=lambda x: x[1], reverse=True)
    return [doc for doc, _ in ranked_results[:top_k]]
 
# def retrieve_documents(prompt):
#     collection_name = "test_files"
#     chroma_client = get_chroma_client()
#     try:
#         collection = chroma_client.get_collection(name=collection_name)
#     except chromadb.errors.InvalidCollectionException:
#         return ""
 
#     initial_results, embeddings = semantic_search(prompt, collection)
#     reranked_results = re_rank(prompt, initial_results, embeddings)
   
#     context = " ".join(reranked_results)
#     return context
 
def get_llm(prompt):
    if not prompt or not isinstance(prompt, str):
        return "Error: Invalid prompt"
   
    conn = http.client.HTTPSConnection("api.generative.engine.capgemini.com")
    session_id = str(uuid.uuid4())
    files = []
    context = retrieve_documents(prompt)
    print(context)
   
   
    payload = json.dumps({
        "action": "run",
        "modelInterface": "multimodal",
        "data": {
            "mode": "chain",
            "text": prompt,
            "files": [],
            "modelName": "anthropic.claude-3-5-sonnet-20240620-v1:0",
            "provider": "bedrock",
            "systemPrompt": f"""You are an AI chatbot designed specifically for Capgemini Engineering's Wiring Harness domain. Your primary purpose is to assist users in retrieving data based on the provided documents. You must adhere strictly to the context and information within the uploaded documents and workspace.
                    Refer the context to generate the response: \n\n{context}
 
                    1.l the values should refer from context only.
                    2.Do not provide any information that is not present in the context.
                    3.While giving response for the user's query, do not mention everytime as "Based on the provided context..., Acoording to the provided context..."like that.
                    4.Do not mention anything about the provided context,just give the response only based on the context.
                    5.After calculating the derated current value, suggest a part number of a wrie that matches or taht is close to the derated current value.      
                    """,
            "sessionId": session_id,
            "modelKwargs": {
                "maxTokens": 1200,
                "temperature": 0.6,
                "streaming": True,
                "topP": 0.9
            }
        }
    })
   
    headers = {
        'Content-Type': 'application/json',
        'x-api-key': 'ASmRfgQH4s270dHcFBrmh4349TPCqygfaAcmsYML'
    }
   
    conn.request("POST", "/v2/llm/invoke", payload, headers)
    res = conn.getresponse()
    data = res.read()
    response_dict = json.loads(data.decode("utf-8"))
   
    if 'content' in response_dict:
        return response_dict['content']
    elif 'error' in response_dict:
        return f"API Error: {response_dict['error']}"
    else:
        return f"Unexpected API response structure: {response_dict}"
 
def read_files(folder_path):
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if filename.endswith('.pdf'):
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    text = ''.join(page.extract_text() for page in pdf_reader.pages)
                    documents.append({'text': text, 'page': 0})
            elif filename.endswith('.xlsx'):
                excel_data = pd.read_excel(file_path, sheet_name=None)
                text = "\n".join(df.to_string() for df in excel_data.values())
                documents.append({'text': text, 'page': 0})
            elif filename.endswith('.csv'):
                csv_text = read_csv_file(file_path)
                if csv_text:
                    documents.append(csv_text)
            elif filename.endswith('.docx'):
                tables_data = extract_tables_from_docx(file_path)
                for table_data in tables_data:
                    documents.append({'text': table_data, 'page': 0})
               
                doc = docxDocument(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                documents.append({'text': text, 'page': 0})
 
        except Exception as e:
            print(f"Error reading {filename}: {e}")
    return documents
 
def chunk_documents(documents, chunk_size=300):
    chunks = []
    for document in documents:
        if isinstance(document, str):
            if document.startswith("Table"):
                chunks.append(document)
            else:
                words = document.split()
                for i in range(0, len(words), chunk_size):
                    chunk = ' '.join(words[i:i + chunk_size])
                    chunks.append(chunk)
        elif isinstance(document, dict) and 'text' in document:
            if document['text'].startswith("Table"):
                chunks.append(document['text'])
            else:
                words = document['text'].split()
                for i in range(0, len(words), chunk_size):
                    chunk = ' '.join(words[i:i + chunk_size])
                    chunks.append(chunk)
    return chunks
 
def filter_relevant_documents(documents, query_embedding, threshold=0.5):
    relevant_docs = []
    model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
    for doc in documents:
        doc_embedding = model.encode(doc).tolist()
        similarity = cosine_similarity([query_embedding], [doc_embedding])[0][0]
        if similarity > threshold:
            relevant_docs.append(doc)
    return relevant_docs
 
#Main execution
conversation=[]
def app():
    st.title("Wire Harness Chatbot")
 
    # Add a file uploader to the app
    # st.write("Please upload your files below:")
    folder_path = r"C:\Users\barsrini\OneDrive - Capgemini\Chatbot_ChromaDB\Documents"
    # collection = initialize_chromadb(folder_path)
    documents = read_files(folder_path)
    chunks = chunk_documents(documents)
 
    create_collection("test_files", chunks)
   
 
    st.write("Start your conversation with the chatbot below:")
    if "conversation" not in st.session_state:
        st.session_state.conversation = []
    user_input = st.text_input("Ask Question:")
   
 
    if st.button('Send'):
        if user_input:
            # Append user input to the conversation.
            st.session_state.conversation.append(f"You: {user_input}")
 
            llm_response = get_llm(user_input)
            st.session_state.conversation.append(f"Chatbot: {llm_response}")
 
             # Display the conversation.
            for line in st.session_state.conversation:
                st.write(line)
 
# Run the Streamlit app
if __name__ == "__main__":
    app()
